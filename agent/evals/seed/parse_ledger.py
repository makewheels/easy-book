"""解析教练台账（毕业学员.docx + 学员课程表.xlsx）→ data/ledger_parsed.json。

用法：uv run --with python-docx,openpyxl python parse_ledger.py

清洗约定（与真实台账对账的口径，逐条可追溯）：
- 台账覆盖 2025-11 ~ 2026-08：月份 11/12 → 2025，其余 → 2026
- 时间推断：小时 ≤8 一律 +12（下午场），≥9 视为上午；"5.30"→17:30；时间为推断值，日期忠实
- 学员内条目按台账顺序视为时间单调：日号回落 → 下一个月；日号超出当月天数 → 下一个月
- 孤立时间点（如单独"3点"）= 上一条目同日加课；条目后裸露日号（"26号 27 28"）= 仅日期条目
- 整列无月份：取同块学员显式月份的众数兜底（报告中标注）
- Excel 序列日期（如 45981）按 1899-12-30 偏移精确换算
- 名字标注：N节=明确节数；≥100 的数字=价格；其余 1-30 数字=节数候选
  （考勤数超过标注节数时取考勤数，视为续费未记账；N人不算节数）
- 多人合栏（"魏显双，宋雨"）拆为多学员、共享考勤；5 字连写按 3+2 拆分
"""

from __future__ import annotations

import calendar
import json
import re
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

HERE = Path(__file__).resolve().parent
DOCX = HERE / "data" / "毕业学员.docx"
XLSX = HERE / "data" / "学员课程表.xlsx"
OUT = HERE / "data" / "ledger_parsed.json"

CN_NUM = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6,
          "七": 7, "八": 8, "九": 9, "十": 10, "十一": 11, "十二": 12}

ENTRY_RE = re.compile(
    r"(?:(\d{1,2}|十一|十二|十|[一二三四五六七八九])\s*月)?"
    r"(\d{1,2})\s*[号日个]"
    r"(?:\s*晚上)?"
    # 小时必须紧贴日号（"26号 27"的 27 是下一天，不是 27 点）——
    # 隔空的小时（"28号，8点"）落到 TIME_ONLY 同日加课路径，合并结果一致
    # 分钟必须紧贴"点"；后接 月/号/日 不吞；"点"本身必须消费（防假名字/幽灵时间点）
    r"(?:(\d{1,2})(?:[.:：](\d{1,2}))?(?:点(?:(\d{1,2})(?!\s*[月号日]))?)?)?"
)
TIME_ONLY_RE = re.compile(r"(晚上)?\s*(?<![\d.])(\d{1,2})(?:\s*[.:：]\s*(\d{1,2}))?\s*点")
RANGE_ONLY_RE = re.compile(r"(?<![\d.])(\d{1,2})\s*[-~]\s*(\d{1,2})\s*点?")
BARE_DAY_RE = re.compile(r"(?<![\d.])(\d{1,2})(?![\d号日点年月])")
SERIAL_RE = re.compile(r"^4\d{4}$")
PRICE_RE = re.compile(r"(?<![\d.])(\d{3,4})(?![\d.])")


def normalize(text: str) -> str:
    text = str(text).replace("\n", " ").replace("\t", " ")
    # 保留 "." 和 "："：它们是时间记法的一部分（5.30 / 5：30）
    for ch in "，。、；！？@#（）()[]【】\"“”":
        text = text.replace(ch, " ")
    text = re.sub(r"(\d)\s*个\s*(\d)", r"\1号\2", text)  # "12个7点"→"12号7点"
    for cn, num in sorted(CN_NUM.items(), key=lambda kv: -len(kv[0])):
        text = text.replace(f"{cn}月", f"{num}月")
    return text


def decode_time(hour: int | None, minute: int | None) -> str | None:
    """教练记法 → 24h。小时 ≤8 视为下午场 +12；9-11 上午；≥12 原样；越界丢弃。"""
    if hour is None or hour > 23 or (minute or 0) > 59:
        return None
    if hour <= 8:
        hour += 12
    return f"{hour:02d}:{minute or 0:02d}"


def fraction_minute(tok: str | None) -> int | None:
    if tok is None:
        return None
    return int(tok) * 10 if len(tok) == 1 else int(tok)  # "6.3" → 6:30


def extract_entries(text: str, warnings: list) -> tuple[list[dict], str]:  # noqa: C901
    """抽取考勤条目。返回 (entries, 归一化文本)；entry 位置基于归一化文本。"""
    text = normalize(text)
    entries: list[dict] = []
    mask = [False] * len(text)

    for m in ENTRY_RE.finditer(text):
        month = CN_NUM.get(m.group(1)) or (int(m.group(1)) if m.group(1) else None)
        hour = int(m.group(3)) if m.group(3) else None
        minute = fraction_minute(m.group(4))
        if m.group(5):
            minute = int(m.group(5))
        note = None
        if hour is not None:
            tail = re.match(r"\s*[-~]\s*(\d{1,2})", text[m.end(3):m.end(3) + 4])
            if tail:
                note = f"原记录{hour}-{tail.group(1)}点区间"
        entries.append({
            "raw": m.group(0).strip(), "month": month, "day": int(m.group(2)),
            "time": decode_time(hour, minute),
            "time_inferred": hour is not None and hour <= 8,
            "notes": [note] if note else [], "_pos": m.start(), "_end": m.end(),
        })
        for i in range(m.start(), m.end()):
            mask[i] = True

    def masked_text() -> str:
        return "".join(" " if masked else ch for ch, masked in zip(text, mask))

    # 孤立时间点 = 上一条目同日加课
    for m in TIME_ONLY_RE.finditer(masked_text()):
        hour = int(m.group(2))
        if not 1 <= hour <= 23:
            continue
        entries.append({
            "raw": m.group(0).strip(), "month": None, "day": None,
            "time": decode_time(hour, fraction_minute(m.group(3))),
            "time_inferred": hour <= 8, "same_day_extra": True,
            # _pos 取数字本身的位置——前导 \s* 可能横跨被掩码的条目，不能算进来
            "notes": [], "_pos": m.start(2), "_end": m.end(),
        })

    # 孤立区间时间（"30号，3-5"→15:00，记区间备注）
    for m in RANGE_ONLY_RE.finditer(masked_text()):
        hour = int(m.group(1))
        if 1 <= hour <= 23:
            entries.append({
                "raw": m.group(0).strip(), "month": None, "day": None,
                "time": decode_time(hour, 0),
                "time_inferred": hour <= 8, "same_day_extra": True,
                "notes": [f"原记录{hour}-{m.group(2)}点区间"],
                "_pos": m.start(1), "_end": m.end(),
            })

    # 裸露日号（"26号 27 28"）：需同时满足 ① 紧邻已识别条目 ② 与最近条目的日号差 ≤2
    #   （连续日期才算延续；"第3节 16号4点"里夹的课程序号 3 与前一天的日号差太远，拒绝）
    scan = masked_text()
    accepted: set[int] = set()
    day_marks: list[tuple[int, int]] = sorted((e["_pos"], e["day"]) for e in entries if e.get("day"))
    for m in BARE_DAY_RE.finditer(scan):
        day = int(m.group(1))
        j = m.start() - 1
        # 回退跳过普通空格；遇到掩码字符（条目本身）立即停——掩码区也是空格，不能穿过
        while j >= 0 and scan[j] == " " and not mask[j]:
            j -= 1
        prev_day = next((d for p, d in reversed(day_marks) if p < m.start()), None)
        if (j >= 0 and (mask[j] or j in accepted) and 1 <= day <= 31
                and prev_day is not None and abs(day - prev_day) <= 2):
            entries.append({
                "raw": m.group(0), "month": None, "day": day,
                "time": None, "time_inferred": False,
                "notes": ["仅日期，时间未记录"], "_pos": m.start(), "_end": m.end(),
            })
            accepted.add(m.end() - 1)
            day_marks.append((m.start(), day))

    entries.sort(key=lambda e: e["_pos"])
    return entries, text


def find_names(text: str, entries: list[dict]) -> list[tuple[str, int]]:
    """遮掉条目后找 ≥2 字中文名（返回 名字,位置；位置基于归一化文本）。"""
    mask = [False] * len(text)
    for e in entries:
        for i in range(e["_pos"], min(e["_pos"] + len(e["raw"]), len(text))):
            mask[i] = True
    masked = "".join(" " if m else ch for ch, m in zip(text, mask))
    names: list[tuple[str, int]] = []
    for m in re.finditer(r"[\u4e00-\u9fa5]{2,5}", masked):
        run = m.group(0)
        if run in ("晚上", "早上", "上午", "下午") or run[0] == "第" or run.endswith("节"):
            continue
        if len(run) == 5:  # "周洪洲马钰" 型连写按 3+2 拆
            names.append((run[:3], m.start()))
            names.append((run[3:], m.start() + 3))
        else:
            names.append((run, m.start()))
    return names


def mask_entries(text: str, entries: list[dict]) -> str:
    """把条目区间替换为空格——标注解析只看条目之外的数字，防把日号/月份误认成节数。"""
    chars = list(text)
    for e in entries:
        for i in range(e["_pos"], min(e.get("_end", e["_pos"] + len(e["raw"])), len(text))):
            chars[i] = " "
    return "".join(chars)


def parse_annotation(text: str, name: str) -> dict:  # noqa: C901
    """名字前后数字标注 → count/price/notes。≥100 为价格；"N节"明确节数；N人不算节数。"""
    ann = {"count": None, "count_explicit": False, "price": None, "notes": []}
    before, _, after = text.partition(name)

    m = re.search(r"(\d{1,2})\s*节", after[:12])
    if m:
        ann["count"] = int(m.group(1))
        ann["count_explicit"] = True
    pm = PRICE_RE.search(after[:14])
    price_span = pm.span() if pm else None
    if pm:
        ann["price"] = int(pm.group(1))
    if not ann["count"]:
        for m2 in re.finditer(r"\d{1,2}", after[:12]):
            if price_span and price_span[0] <= m2.start() < price_span[1]:
                continue
            if re.match(r"\d+\s*人", after[m2.start():]):
                ann["notes"].append("多人同行")
                break
            if 1 <= int(m2.group(0)) <= 30:
                ann["count"] = int(m2.group(0))
                break
    if not ann["count"]:
        lead = re.search(r"(\d{1,2})\s*$", before)
        if lead:
            ann["count"] = int(lead.group(1))
            ann["count_leading"] = True
    for mark in ("蛙", "哈"):
        if mark in after[:8]:
            ann["notes"].append(f"标注:{mark}")
    if "+自" in text or re.search(r"自\s*$", after[:6]):
        ann["notes"].append("标注:+自")
    return ann


# ── 日期序列化：年份推断 + 月份补全 ─────────────────────────
def year_for_month(month: int) -> int:
    return 2025 if month >= 11 else 2026


def resolve_dates(entries: list[dict], warnings: list, name: str,  # noqa: C901, PLR0912, PLR0915
                  initial_month: int | None = None) -> None:
    # 孤立时间点挂到前一条：前一条没时间 → 补全（"11月12日\t3点"是一节课）；
    # 前一条已有时间 → 同日加课（"5月6号2点 / 3点"是两节）
    merged: list[dict] = []
    for e in entries:
        if e.get("same_day_extra"):
            if not merged:
                warnings.append(f"{name}: 孤立时间点无前置日期: {e['raw']}")
            elif merged[-1].get("time") is None:
                merged[-1]["time"] = e["time"]
                merged[-1]["time_inferred"] = e["time_inferred"]
                merged[-1]["raw"] = f"{merged[-1]['raw']} {e['raw']}"
                merged[-1]["notes"] = merged[-1].get("notes", []) + e["notes"]
            else:
                clone = dict(merged[-1])
                clone.update({"raw": e["raw"], "time": e["time"],
                              "time_inferred": e["time_inferred"], "notes": e["notes"]})
                merged.append(clone)
        else:
            merged.append(e)
    entries[:] = merged
    if not entries:
        return

    explicit_idx = [i for i, e in enumerate(entries) if e.get("month")]
    if not explicit_idx:
        if initial_month:
            entries[0]["month"] = initial_month
            entries[0]["notes"].append(f"月份取块内众数{initial_month}月")
            warnings.append(f"{name}: 无显式月份，按块内众数取{initial_month}月")
            explicit_idx = [0]
        else:
            warnings.append(f"{name}: 全部条目无月份，无法定位日期（共{len(entries)}条）")
            return

    def days_in(m: int) -> int:
        return calendar.monthrange(year_for_month(m), m)[1]

    anchor = explicit_idx[0]
    months: dict[int, int] = {anchor: entries[anchor]["month"]}

    cur_m, cur_day = entries[anchor]["month"], entries[anchor]["day"]
    for i in range(anchor + 1, len(entries)):
        e = entries[i]
        if e.get("month"):
            cur_m = e["month"]
        else:
            if e["day"] < cur_day:
                cur_m = cur_m % 12 + 1
            if e["day"] > days_in(cur_m):  # "6月31号" → 顺延下月并继续传播
                cur_m = cur_m % 12 + 1
        months[i] = cur_m
        cur_day = e["day"]

    cur_m, cur_day = entries[anchor]["month"], entries[anchor]["day"]
    for i in range(anchor - 1, -1, -1):
        e = entries[i]
        if e.get("month"):
            cur_m = e["month"]
        else:
            if e["day"] >= cur_day:
                cur_m = (cur_m - 2) % 12 + 1
            if e["day"] > days_in(cur_m):
                cur_m = (cur_m - 2) % 12 + 1
        months[i] = cur_m
        cur_day = e["day"]

    prev_date = None
    for i, e in enumerate(entries):
        m = months[i]
        day = e["day"]
        if day > days_in(m):  # 游走兜不住（如 2月29）→ 钳到月末
            day = days_in(m)
            warnings.append(f"{name}: {m}月{e['day']}号不存在，钳到月末{day}号")
        try:
            d = date(year_for_month(m), m, day)
        except ValueError:
            warnings.append(f"{name}: 非法日期 {m}月{e['day']}号（{e['raw']}），丢弃")
            e["date"] = None
            continue
        if prev_date and d < prev_date:
            warnings.append(f"{name}: 日期回退 {prev_date}→{d}（{e['raw']}），按台账顺序保留")
        e["date"] = d.isoformat()
        prev_date = d


def build_attendance(st: dict, warnings: list, initial_month: int | None = None) -> None:
    resolve_dates(st["_raw_entries"], warnings, st["name"], initial_month)
    att = [{
        "raw": e["raw"], "date": e.get("date"), "time": e.get("time"),
        "time_inferred": e.get("time_inferred", False), "notes": e.get("notes", []),
    } for e in st["_raw_entries"] if e.get("date")]
    att.sort(key=lambda a: (a["date"], a["time"] or ""))
    deduped, seen = [], set()
    for a in att:
        key = (a["date"], a["time"])
        if key not in seen:
            seen.add(key)
            deduped.append(a)
    st["attendance"] = deduped
    st.pop("_raw_entries", None)


def finish_blocks(blocks: list[list[dict]], warnings: list) -> None:
    """块内众数月份兜底整列无月份的学员，然后构建 attendance。"""
    for block in blocks:
        counts: dict[int, int] = {}
        for st in block:
            for e in st["_raw_entries"]:
                if e.get("month"):
                    counts[e["month"]] = counts.get(e["month"], 0) + 1
        mode = max(counts, key=counts.get) if counts else None
        for st in block:
            has_month = any(e.get("month") for e in st["_raw_entries"])
            build_attendance(st, warnings, None if has_month else mode)


def new_student(name: str, ann: dict, sources: list, graduated: bool) -> dict:
    return {
        "name": name, "graduated": graduated, "sources": sources,
        "package": {"count": ann["count"], "count_explicit": ann.get("count_explicit", False),
                    "price": ann["price"], "notes": ann["notes"]},
        "attendance": [], "_raw_entries": [],
    }


# ── docx 解析 ───────────────────────────────────────────────
def parse_docx(warnings: list) -> tuple[list[dict], list[list[dict]]]:  # noqa: C901, PLR0912
    doc = Document(DOCX)
    students: list[dict] = []
    blocks: list[list[dict]] = []

    # 段落区：名字段 + 条目段交替
    current: list[dict] = []
    para_block: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        entries, ntext = extract_entries(text, warnings)
        names = find_names(ntext, entries)
        if names:
            masked = mask_entries(ntext, entries)
            current = [new_student(nm, parse_annotation(masked, nm),
                                   ["毕业学员.docx"], True) for nm, _ in names]
            para_block.extend(current)
            students.extend(current)
        else:
            for st in current:
                st["_raw_entries"].extend(entries)
    blocks.append(para_block)

    # 表格区：单元格内多学员按名字切分；无主条目单元格跟随表格内当前学员
    seen_cells: set[str] = set()
    for table in doc.tables:
        table_block: list[dict] = []
        table_current: list[dict] = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text or text in seen_cells:  # 合并单元格会重复读出
                    continue
                seen_cells.add(text)
                entries, ntext = extract_entries(text, warnings)
                names = find_names(ntext, entries)
                if not names:
                    if entries and table_current:
                        for st in table_current:
                            st["_raw_entries"].extend(entries)
                    elif entries:
                        warnings.append(f"docx 表格有无主考勤条目: {ntext[:40]}")
                    continue
                positions = sorted([(pos, nm) for nm, pos in names])
                masked = mask_entries(ntext, entries)
                table_current = []
                for idx, (pos, nm) in enumerate(positions):
                    # 段起点取上一个名字末尾——名字前的数字标注（"10张晶晶1700"）不能丢
                    seg_start = positions[idx - 1][0] + len(positions[idx - 1][1]) if idx else 0
                    end = positions[idx + 1][0] if idx + 1 < len(positions) else len(ntext)
                    seg = masked[seg_start:end]
                    st = new_student(nm, parse_annotation(seg, nm), ["毕业学员.docx"], True)
                    st["_raw_entries"] = [e for e in entries if pos <= e["_pos"] < end]
                    table_current.append(st)
                    table_block.append(st)
                    students.append(st)
        blocks.append(table_block)
    return students, blocks


# ── xlsx 解析 ───────────────────────────────────────────────
def parse_xlsx(warnings: list) -> tuple[list[dict], list[list[dict]]]:  # noqa: C901, PLR0912
    wb = load_workbook(XLSX, data_only=True)
    students: list[dict] = []
    blocks: list[list[dict]] = []
    for ws in wb.worksheets:
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
        header_rows: list[tuple[int, dict[int, list[dict]]]] = []
        for ri, row in enumerate(rows):
            header_cells: dict[int, list[dict]] = {}
            for ci, val in enumerate(row):
                if val is None or not str(val).strip():
                    continue
                entries, ntext = extract_entries(str(val), warnings)
                names = find_names(ntext, entries)
                if names:
                    masked = mask_entries(ntext, entries)
                    header_cells[ci] = [
                        new_student(nm, parse_annotation(masked, nm),
                                    ["学员课程表.xlsx"], False)
                        for nm, _ in names
                    ]
            if len(header_cells) >= 3:
                header_rows.append((ri, header_cells))

        for hi, (hrow, cols) in enumerate(header_rows):
            block_students = [st for sts in cols.values() for st in sts]
            blocks.append(block_students)
            students.extend(block_students)
            stop = header_rows[hi + 1][0] if hi + 1 < len(header_rows) else len(rows)
            for ci, sts in cols.items():
                for ri in range(hrow + 1, stop):
                    val = rows[ri][ci] if ci < len(rows[ri]) else None
                    if val is None or not str(val).strip():
                        break
                    text = str(val).strip()
                    if SERIAL_RE.match(text):  # Excel 序列日期
                        d = (datetime(1899, 12, 30) + timedelta(days=int(text))).date()
                        for st in sts:
                            st["_raw_entries"].append({
                                "raw": text, "month": d.month, "day": d.day,
                                "time": None, "time_inferred": False,
                                "notes": ["Excel序列日期，时间未记录"], "_pos": 0,
                            })
                        continue
                    entries, _ = extract_entries(text, warnings)
                    if not entries:
                        warnings.append(f"xlsx 疑似补记（无考勤时间）: {text}")
                        continue
                    for st in sts:
                        st["_raw_entries"].extend(entries)
    return students, blocks


def merge_duplicates(students: list[dict], warnings: list) -> list[dict]:
    merged: dict[str, dict] = {}
    order: list[str] = []
    for st in students:
        if st["name"] in merged:
            base = merged[st["name"]]
            for src in st["sources"]:
                if src not in base["sources"]:
                    base["sources"].append(src)
            base["graduated"] = base["graduated"] and st["graduated"]
            base["attendance"].extend(st["attendance"])
            for k in ("count", "price"):
                if base["package"].get(k) is None:
                    base["package"][k] = st["package"].get(k)
            warnings.append(f"{st['name']}: 跨文件合并")
        else:
            merged[st["name"]] = st
            order.append(st["name"])
    result = []
    for name in order:
        st = merged[name]
        att, seen = [], set()
        for a in sorted(st["attendance"], key=lambda x: (x["date"], x["time"] or "")):
            key = (a["date"], a["time"])
            if key not in seen:
                seen.add(key)
                att.append(a)
        st["attendance"] = att
        result.append(st)
    return result


def main() -> None:
    warnings: list[str] = []
    docx_students, docx_blocks = parse_docx(warnings)
    xlsx_students, xlsx_blocks = parse_xlsx(warnings)
    finish_blocks(docx_blocks, warnings)
    finish_blocks(xlsx_blocks, warnings)
    all_students = merge_duplicates(docx_students + xlsx_students, warnings)

    total_att = sum(len(s["attendance"]) for s in all_students)
    no_att = [s["name"] for s in all_students if not s["attendance"]]
    result = {
        "source_files": ["毕业学员.docx", "学员课程表.xlsx"],
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "conventions": {
            "year": "11/12月→2025，其余→2026（台账覆盖 2025-11~2026-08）",
            "time": "小时≤8按下午场+12；9-11上午；≥12原样；5.30→17:30；时间为推断值，日期忠实",
            "count": "考勤数>标注节数时取考勤数（视为续费未记账）；N人不算节数",
            "venue_share": "台账无分成记录，导入统一 0，后续可用 agent 设置",
            "multi_person": "合栏学员共享考勤；赵家三口按单一家庭记录",
        },
        "students": all_students,
        "warnings": warnings,
        "stats": {
            "students": len(all_students),
            "graduated": sum(1 for s in all_students if s["graduated"]),
            "attendance_records": total_att,
            "with_price": sum(1 for s in all_students if s["package"]["price"]),
            "with_count": sum(1 for s in all_students if s["package"]["count"]),
            "no_attendance": no_att,
        },
    }
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    s = result["stats"]
    print(f"学员 {s['students']}（毕业 {s['graduated']}），考勤 {s['attendance_records']} 条")
    print(f"有价格 {s['with_price']}，有节数 {s['with_count']}，无考勤 {len(s['no_attendance'])}: {s['no_attendance']}")
    print(f"警告 {len(warnings)} 条：")
    for w in warnings:
        print("  -", w)


if __name__ == "__main__":
    main()
